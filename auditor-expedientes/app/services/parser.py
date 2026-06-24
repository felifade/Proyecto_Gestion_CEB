import fitz  # PyMuPDF
import pdfplumber
import pandas as pd
import docx  # python-docx
import pytesseract
from PIL import Image
import io
import os

def extract_text_from_pdf(file_path: str, file_data: bytes = None) -> tuple[str, int]:
    """
    Extracts text from a PDF file using PyMuPDF (fitz), with a fallback to pdfplumber.
    If no text is extracted (scanned document), it uses pytesseract OCR on each page.
    Supports in-memory bytes or path-based reading.
    """
    text = ""
    total_pages = 0
    try:
        # 1. Try PyMuPDF (fastest)
        if file_data is not None:
            doc = fitz.open(stream=file_data, filetype="pdf")
        else:
            doc = fitz.open(file_path)
            
        total_pages = len(doc)
        for page in doc:
            text += page.get_text()
        doc.close()
        
        # 2. Fallback to pdfplumber if text is minimal
        if len(text.strip()) < 50:
            text = ""
            pdf_src = io.BytesIO(file_data) if file_data is not None else file_path
            with pdfplumber.open(pdf_src) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    
        # 3. Fallback to pytesseract OCR if still empty (image-only PDF)
        if len(text.strip()) < 50:
            text = ""
            if file_data is not None:
                doc = fitz.open(stream=file_data, filetype="pdf")
            else:
                doc = fitz.open(file_path)
                
            # Limit OCR to first 3 pages to optimize speed and prevent cloud timeouts
            max_ocr_pages = min(len(doc), 3)
            for page_num in range(max_ocr_pages):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                image_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(image_data))
                # Run OCR in Spanish
                page_text = pytesseract.image_to_string(image, lang='spa')
                text += f"\n--- Página {page_num + 1} ---\n" + page_text
            doc.close()
            
    except Exception as e:
        text = f"Error al extraer texto del PDF: {str(e)}"
    return text, total_pages

def extract_text_from_excel(file_path: str, file_data: bytes = None) -> str:
    """
    Extracts content from all sheets in an Excel file using pandas.
    """
    try:
        excel_src = io.BytesIO(file_data) if file_data is not None else file_path
        xls = pd.ExcelFile(excel_src)
        sheet_texts = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            sheet_texts.append(f"Hoja: {sheet_name}\n" + df.to_string(index=False))
        return "\n\n".join(sheet_texts)
    except Exception as e:
        return f"Error al extraer texto de Excel: {str(e)}"

def extract_text_from_docx(file_path: str, file_data: bytes = None) -> str:
    """
    Extracts text from paragraphs and tables in a Word document.
    """
    try:
        docx_src = io.BytesIO(file_data) if file_data is not None else file_path
        doc = docx.Document(docx_src)
        full_text = []
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        return f"Error al extraer texto de Word: {str(e)}"

def parse_document(file_path: str, file_data: bytes = None) -> tuple[str, int]:
    """
    Determines document type and extracts text. Returns (text, page_count).
    Supports in-memory bytes parsing if file_data is supplied.
    """
    ext = file_path.split(".")[-1].lower() if "." in file_path else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_path, file_data)
    elif ext in ["xlsx", "xls"]:
        return extract_text_from_excel(file_path, file_data), 1
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_path, file_data), 1
    else:
        try:
            if file_data is not None:
                content = file_data.decode("utf-8", errors="ignore")[:10000]
                return content, 1
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(10000) # Read first 10k chars
                    return content, 1
        except Exception:
            return "", 0
