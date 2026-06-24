import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from sqlalchemy.orm import Session
from ..models import Expediente, ResultadoAuditoria, Criteria

def generate_consolidated_excel(db: Session, output_path: str):
    """
    Generates a consolidated spreadsheet listing all audited folders (rows) 
    and criteria evaluations (columns) formatted with status colors.
    """
    expedientes = db.query(Expediente).all()
    criterios = db.query(Criteria).all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Consolidado Auditoría"
    
    # 1. Header Row
    headers = ["Carpeta Expediente", "Año", "Estado Global", "% Cumplimiento"] + [c.criterio for c in criterios]
    ws.append(headers)
    
    # Header styles
    fill_header = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid") # Slate-800
    font_header = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="center")
    
    # Apply header style
    ws.row_dimensions[1].height = 28
    for col in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col)
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = align_center
        
    # Cell fills for status colors (soft palette)
    fill_green = PatternFill(start_color="D1FAE5", end_color="D1FAE5", fill_type="solid")  # soft green
    fill_yellow = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid") # soft yellow
    fill_red = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")    # soft red
    fill_gray = PatternFill(start_color="F3F4F6", end_color="F3F4F6", fill_type="solid")   # soft gray
    
    font_bold = Font(name="Calibri", size=10, bold=True)
    font_normal = Font(name="Calibri", size=10)
    
    border_thin = Border(
        left=Side(style='thin', color='E2E8F0'),
        right=Side(style='thin', color='E2E8F0'),
        top=Side(style='thin', color='E2E8F0'),
        bottom=Side(style='thin', color='E2E8F0')
    )
    
    # 2. Append Data Rows
    for r_idx, exp in enumerate(expedientes, start=2):
        row = [exp.nombre_carpeta, exp.anio or "General", exp.resultado_global, f"{exp.porcentaje_cumplimiento:.1f}%"]
        
        for crit in criterios:
            res = db.query(ResultadoAuditoria).filter(
                ResultadoAuditoria.expediente_id == exp.id,
                ResultadoAuditoria.criterio_id == crit.id
            ).first()
            row.append(res.estado if res else "Sin analizar")
            
        ws.append(row)
        ws.row_dimensions[r_idx].height = 20
        
        # Style row cells
        ws.cell(row=r_idx, column=1).font = font_bold
        ws.cell(row=r_idx, column=1).alignment = align_left
        ws.cell(row=r_idx, column=1).border = border_thin
        
        # Año
        ws.cell(row=r_idx, column=2).font = font_normal
        ws.cell(row=r_idx, column=2).alignment = align_center
        ws.cell(row=r_idx, column=2).border = border_thin
        
        # Global Result styling
        cell_res = ws.cell(row=r_idx, column=3)
        cell_res.font = font_bold
        cell_res.alignment = align_center
        cell_res.border = border_thin
        if exp.resultado_global == "Cumple":
            cell_res.fill = fill_green
        elif exp.resultado_global == "Cumple parcialmente":
            cell_res.fill = fill_yellow
        elif exp.resultado_global == "No cumple":
            cell_res.fill = fill_red
        else:
            cell_res.fill = fill_gray
            
        # Percentage styling
        cell_pct = ws.cell(row=r_idx, column=4)
        cell_pct.font = font_bold
        cell_pct.alignment = align_center
        cell_pct.border = border_thin
        
        # Criteria columns styling
        for c_idx, crit in enumerate(criterios, start=5):
            cell_crit = ws.cell(row=r_idx, column=c_idx)
            cell_crit.font = font_normal
            cell_crit.alignment = align_center
            cell_crit.border = border_thin
            status = cell_crit.value
            
            if status == "Cumple":
                cell_crit.fill = fill_green
            elif status == "Cumple parcialmente":
                cell_crit.fill = fill_yellow
            elif status == "No cumple":
                cell_crit.fill = fill_red
            else:
                cell_crit.fill = fill_gray
                
    # Auto-adjust column widths
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = col[0].column_letter
        ws.column_dimensions[col_letter].width = max(max_len + 3, 12)
        
    wb.save(output_path)

def generate_executive_word(db: Session, output_path: str):
    """
    Generates a formal Microsoft Word (.docx) executive report with charts
    descriptions, metadata summary, and audit observations.
    """
    doc = Document()
    
    # 1. Title Section
    title = doc.add_paragraph()
    title_run = title.add_run("REPORTE EJECUTIVO DE AUDITORÍA")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = 'Arial'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"Generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}")
    sub_run.italic = True
    sub_run.font.size = Pt(11)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run("\n") # Spacer
    
    # 2. General Statistics Summary
    doc.add_heading("1. Resumen de Estadísticas Generales", level=1)
    
    expedientes = db.query(Expediente).all()
    total = len(expedientes)
    cumplen = sum(1 for e in expedientes if e.resultado_global == "Cumple")
    parciales = sum(1 for e in expedientes if e.resultado_global == "Cumple parcialmente")
    no_cumplen = sum(1 for e in expedientes if e.resultado_global == "No cumple")
    pendientes = sum(1 for e in expedientes if e.estado_analisis == "Pendiente")
    
    p = doc.add_paragraph()
    p.add_run(f"• Total de expedientes detectados: ").bold = True
    p.add_run(f"{total}\n")
    p.add_run(f"• Expedientes que cumplen al 100%: ").bold = True
    p.add_run(f"{cumplen}\n")
    p.add_run(f"• Expedientes con cumplimiento parcial: ").bold = True
    p.add_run(f"{parciales}\n")
    p.add_run(f"• Expedientes con incumplimiento: ").bold = True
    p.add_run(f"{no_cumplen}\n")
    p.add_run(f"• Expedientes pendientes de análisis: ").bold = True
    p.add_run(f"{pendientes}\n")
    
    # 3. Detailed Audit
    doc.add_heading("2. Desglose Detallado por Expediente", level=1)
    
    for exp in expedientes:
        year_str = f" ({exp.anio})" if exp.anio else ""
        doc.add_heading(f"Carpeta: {exp.nombre_carpeta}{year_str}", level=2)
        doc.add_paragraph(f"Porcentaje de Cumplimiento Ponderado: {exp.porcentaje_cumplimiento:.1f}% ({exp.resultado_global})")
        
        resultados = db.query(ResultadoAuditoria).filter(ResultadoAuditoria.expediente_id == exp.id).all()
        
        if not resultados:
            doc.add_paragraph("Este expediente no cuenta con análisis finalizado o fue omitido.")
            continue
            
        # Create Table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Shading Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Criterio de Cotejo'
        hdr_cells[1].text = 'Estado'
        hdr_cells[2].text = 'Observación de la Auditoría'
        
        for r in resultados:
            row_cells = table.add_row().cells
            row_cells[0].text = r.criterio.criterio if r.criterio else "General"
            row_cells[1].text = r.estado
            row_cells[2].text = r.observacion or "Sin observaciones registradas."
            
            # Show Evidences if they exist
            if r.evidencia_texto:
                row_cells[2].text += f"\n[Evidencia: {r.evidencia_documento} (Pág. {r.evidencia_pagina}): \"{r.evidencia_texto}\"]"
                
        doc.add_paragraph().add_run("\n") # Spacer
        
    doc.save(output_path)

def generate_executive_pdf(db: Session, output_path: str):
    """
    Generates a high-quality stylized PDF file summarizing audit results.
    """
    # Initialize Document margins
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    
    styles = getSampleStyleSheet()
    
    # Custom colors and styles
    color_primary = colors.HexColor("#1E293B")
    color_secondary = colors.HexColor("#4F46E5")
    
    style_title = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=color_primary,
        alignment=1, # Center
        spaceAfter=15
    )
    
    style_heading = ParagraphStyle(
        name='HeadingStyle',
        parent=styles['Heading2'],
        fontSize=12,
        leading=16,
        textColor=color_secondary,
        spaceBefore=10,
        spaceAfter=6
    )
    
    style_text = ParagraphStyle(
        name='TextStyle',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        spaceAfter=4
    )
    
    style_table_header = ParagraphStyle(
        name='TableHeader',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        bold=True,
        textColor=colors.white
    )
    
    story = []
    
    # Header Title
    story.append(Paragraph("<b>REPORTE DE AUDITORÍA DE EXPEDIENTES</b>", style_title))
    story.append(Paragraph(f"Fecha de Reporte: {datetime.now().strftime('%d/%m/%Y a las %H:%M')}", style_text))
    story.append(Spacer(1, 15))
    
    expedientes = db.query(Expediente).all()
    
    for exp in expedientes:
        year_str = f" ({exp.anio})" if exp.anio else ""
        story.append(Paragraph(f"<b>Expediente:</b> {exp.nombre_carpeta}{year_str}", style_heading))
        story.append(Paragraph(f"Cumplimiento: <b>{exp.porcentaje_cumplimiento:.1f}%</b> | Resultado: <b>{exp.resultado_global}</b>", style_text))
        
        resultados = db.query(ResultadoAuditoria).filter(ResultadoAuditoria.expediente_id == exp.id).all()
        if not resultados:
            story.append(Paragraph("<i>Pendiente de auditar.</i>", style_text))
            story.append(Spacer(1, 10))
            continue
            
        # Table data
        table_data = [[
            Paragraph("<b>Criterio</b>", style_table_header),
            Paragraph("<b>Estado</b>", style_table_header),
            Paragraph("<b>Observación / Evidencia</b>", style_table_header)
        ]]
        
        for r in resultados:
            obs_text = r.observacion or ""
            if r.evidencia_texto:
                obs_text += f"<br/><font color='#4F46E5'><b>Evidencia:</b> {r.evidencia_documento} (Pág. {r.evidencia_pagina}) - \"{r.evidencia_texto}\"</font>"
                
            table_data.append([
                Paragraph(r.criterio.criterio if r.criterio else "", style_text),
                Paragraph(f"<b>{r.estado}</b>", style_text),
                Paragraph(obs_text, style_text)
            ])
            
        # Build Table
        col_widths = [150, 60, 330]
        t = Table(table_data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), color_primary),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 6),
            ('TOPPADDING', (0,0), (-1,0), 6),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F8FAFC")]),
        ]))
        
        story.append(t)
        story.append(Spacer(1, 15))
        
    doc.build(story)
